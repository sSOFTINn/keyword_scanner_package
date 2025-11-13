# scanner_core.py
import os
import zipfile
import concurrent.futures
from io import BytesIO
from datetime import datetime
import string
import re

from docx import Document as DocxDocument
from docx import Document as ReportDocument
from PyPDF2 import PdfReader
import rarfile
import py7zr


# ===== НАЛАШТУВАННЯ =====

DOCUMENT_EXTENSIONS = {
    ".txt",
    ".docx",
    ".pdf",
}

ARCHIVE_EXTENSIONS = {".zip", ".rar", ".7z"}

SUPPORTED_EXTENSIONS = DOCUMENT_EXTENSIONS | ARCHIVE_EXTENSIONS

DEFAULT_MAX_FILE_SIZE_MB = 200

FAST_EXCLUDED_DIR_NAMES = {
    "winsxs",
    "installer",
    "$recycle.bin",
    "appdata",
    "program files",
    "program files (x86)",
    "programdata",
}


class KeywordScanner:
    """
    mode = "deep"  - повний скан
    mode = "fast"  - швидкий (пропускає деякі системні теки)
    mode = "names" - тільки назви файлів (без читання вмісту)
    """

    def __init__(
        self,
        dictionary_path,
        filename_patterns_path=None,
        mode="deep",                 # "deep" | "fast" | "names"
        max_file_size_mb=DEFAULT_MAX_FILE_SIZE_MB,
        regex_pattern: str | None = None,      # regex по ВМІСТУ
        match_case: bool = False,              # регістр для вмісту
        name_regex_pattern: str | None = None, # regex по НАЗВАХ
        name_match_case: bool = False,         # регістр для назв
    ):
        self.mode = mode.lower()
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024

        # словник (вміст)
        self.dictionary = self.load_dictionary(dictionary_path)
        self.dictionary_lower = [w.lower() for w in self.dictionary]

        # список слів для назв (File List для імен)
        self.filename_patterns = self.load_filename_patterns(filename_patterns_path)

        # параметри пошуку у ВМІСТІ
        self.match_case = match_case
        self.regex_pattern = (regex_pattern or "").strip()
        self.regex = None
        if self.regex_pattern:
            flags = re.MULTILINE | re.DOTALL
            if not self.match_case:
                flags |= re.IGNORECASE
            try:
                self.regex = re.compile(self.regex_pattern, flags)
            except re.error:
                self.regex = None

        # параметри пошуку по НАЗВАХ
        self.name_match_case = name_match_case
        self.name_regex_pattern = (name_regex_pattern or "").strip()
        self.name_regex = None
        if self.name_regex_pattern:
            flags = 0
            if not self.name_match_case:
                flags |= re.IGNORECASE
            try:
                self.name_regex = re.compile(self.name_regex_pattern, flags)
            except re.error:
                self.name_regex = None

        self.matches: list[tuple[str, list[str]]] = []

    # ---------- СЛОВНИКИ / FILE LIST ----------

    def load_dictionary(self, path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return [line.strip() for line in f if line.strip()]
        except Exception:
            return []

    def load_filename_patterns(self, path):
        """
        File List для назв файлів:
        кожен рядок у файлі = окремий шаблон (пошук по входженню в ім'я / шлях).
        """
        if not path:
            return []
        if not os.path.exists(path):
            return []
        try:
            with open(path, "r", encoding="utf-8") as f:
                return [line.strip().lower() for line in f if line.strip()]
        except Exception:
            return []

    # ---------- ЧИТАННЯ ВМІСТУ ----------

    def extract_text_txt(self, path):
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""

    def extract_text_docx(self, path_or_stream):
        try:
            doc = DocxDocument(path_or_stream)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""

    def extract_text_pdf(self, path_or_stream):
        try:
            reader = PdfReader(path_or_stream)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return ""

    # ---------- ПОШУК У НАЗВІ ----------

    def search_in_filename(self, path):
        """
        File List + regex по назві.
        """
        hits: list[str] = []

        # File List (із txt)
        if self.filename_patterns:
            name = os.path.basename(path).lower()
            full = path.lower()
            for pat in self.filename_patterns:
                if pat in name or pat in full:
                    hits.append(f"[NAME]{pat}")

        # regex для назв
        if self.name_regex and self.name_regex.search(path):
            hits.append(f"[NAME_REGEX]{self.name_regex_pattern}")

        return hits

    # ---------- ПОШУК У ВМІСТІ ----------

    def search_in_content(self, text: str):
        hits: list[str] = []
        if not text:
            return hits

        # словникові слова
        if self.match_case:
            src = text
            for w in self.dictionary:
                if w and w in src:
                    hits.append(w)
        else:
            src = text.lower()
            for w in self.dictionary_lower:
                if w and w in src:
                    hits.append(w)

        # regex по вмісту
        if self.regex and self.regex.search(text):
            hits.append(f"[REGEX]{self.regex_pattern}")

        return hits

    # ---------- ВНУТРІШНІ ФАЙЛИ В АРХІВАХ ----------

    def _scan_archive_entry(self, file_stream, ext, parent_path):
        filename_hits = self.search_in_filename(parent_path)

        if self.mode == "names":
            if filename_hits:
                return [(parent_path, filename_hits)]
            return []

        try:
            if ext == ".txt":
                content = file_stream.read().decode("utf-8", errors="ignore")
            elif ext == ".docx":
                content = self.extract_text_docx(BytesIO(file_stream.read()))
            elif ext == ".pdf":
                content = self.extract_text_pdf(BytesIO(file_stream.read()))
            elif ext in ARCHIVE_EXTENSIONS:
                return self.scan_stream_archive(file_stream.read(), ext, parent_path)
            else:
                return []

            content_hits = self.search_in_content(content)
            hits = filename_hits + content_hits
            if hits:
                return [(parent_path, hits)]
            return []
        except Exception:
            return []

    def scan_stream_archive(self, data, ext, parent_path):
        matches = []
        try:
            if ext == ".zip":
                with zipfile.ZipFile(BytesIO(data), "r") as archive:
                    for name in archive.namelist():
                        inner_ext = os.path.splitext(name)[1].lower()
                        if inner_ext in SUPPORTED_EXTENSIONS:
                            with archive.open(name) as file:
                                matches += self._scan_archive_entry(
                                    file, inner_ext, f"{parent_path} -> {name}"
                                )

            elif ext == ".rar":
                with rarfile.RarFile(BytesIO(data)) as archive:
                    for info in archive.infolist():
                        inner_ext = os.path.splitext(info.filename)[1].lower()
                        if inner_ext in SUPPORTED_EXTENSIONS:
                            with archive.open(info) as file:
                                matches += self._scan_archive_entry(
                                    file, inner_ext, f"{parent_path} -> {info.filename}"
                                )

            elif ext == ".7z":
                with py7zr.SevenZipFile(BytesIO(data), mode="r") as archive:
                    for name in archive.getnames():
                        inner_ext = os.path.splitext(name)[1].lower()
                        if inner_ext in SUPPORTED_EXTENSIONS:
                            with archive.read([name])[name] as file:
                                matches += self._scan_archive_entry(
                                    file, inner_ext, f"{parent_path} -> {name}"
                                )
        except Exception:
            pass
        return matches

    # ---------- СКАН ОКРЕМОГО ФАЙЛУ ----------

    def scan_file(self, path):
        ext = os.path.splitext(path)[1].lower()

        name_hits = self.search_in_filename(path)

        if self.mode == "names":
            if name_hits:
                return [(path, name_hits)]
            return []

        if ext not in SUPPORTED_EXTENSIONS:
            if name_hits:
                return [(path, name_hits)]
            return []

        try:
            size = os.path.getsize(path)
            if size > self.max_file_size_bytes:
                if name_hits:
                    return [(path, name_hits)]
                return []
        except Exception:
            pass

        if ext in ARCHIVE_EXTENSIONS:
            try:
                with open(path, "rb") as f:
                    data = f.read()
                return self.scan_stream_archive(data, ext, path)
            except Exception:
                return []

        if ext == ".txt":
            text = self.extract_text_txt(path)
        elif ext == ".docx":
            text = self.extract_text_docx(path)
        elif ext == ".pdf":
            text = self.extract_text_pdf(path)
        else:
            return []

        hits = list(name_hits)
        hits += self.search_in_content(text)

        if hits:
            return [(path, hits)]
        return []

    # ---------- ОБХІД ФС ----------

    def iter_files(self, root):
        stack = [root]
        fast_mode = self.mode == "fast"

        while stack:
            current = stack.pop()
            base = os.path.basename(current).lower()

            if fast_mode and base in FAST_EXCLUDED_DIR_NAMES:
                continue

            try:
                with os.scandir(current) as it:
                    for entry in it:
                        try:
                            if entry.is_dir(follow_symlinks=False):
                                stack.append(entry.path)
                            elif entry.is_file(follow_symlinks=False):
                                yield entry.path
                        except Exception:
                            continue
            except Exception:
                continue

    def scan_directory(self, folder_path, update_callback=None):
        file_paths = list(self.iter_files(folder_path))
        if not file_paths:
            return

        cpu = os.cpu_count() or 4
        workers = min(8, cpu * 2)  # не більше 8 потоків

        with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as executor:
            futures = {executor.submit(self.scan_file, p): p for p in file_paths}
            for future in concurrent.futures.as_completed(futures):
                try:
                    result = future.result()
                except Exception:
                    continue
                if result and update_callback:
                    update_callback(result)

    # ---------- ДИСКИ ----------

    def get_all_drives(self):
        drives = []
        for letter in string.ascii_uppercase:
            d = f"{letter}:\\"
            if os.path.exists(d):
                drives.append(d)
        return drives

    def scan_selected_drives(self, drives, update_callback=None):
        for drive in drives:
            try:
                self.scan_directory(drive, update_callback=update_callback)
            except Exception:
                continue

    def scan_all_drives(self, update_callback=None):
        self.scan_selected_drives(self.get_all_drives(), update_callback=update_callback)

    # ---------- ЗВІТ DOCX ----------

    def generate_report(self, output_path="scan_report.docx"):
        doc = ReportDocument()
        doc.add_heading("Keyword Scan Report", 0)
        doc.add_paragraph(
            f"Scan completed on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        doc.add_paragraph(f"Total matches: {len(self.matches)}\n")

        for path, words in self.matches:
            doc.add_paragraph(f"File: {path}", style="List Bullet")
            doc.add_paragraph(f"Keywords / patterns: {', '.join(words)}")

        doc.save(output_path)
        print(f"Report saved to {output_path}")